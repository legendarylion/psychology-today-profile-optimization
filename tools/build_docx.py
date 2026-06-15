#!/usr/bin/env python3
"""
build_docx.py - turn a validated optimized-profile.md into a client-friendly .docx.

The .docx puts every PT copy block in its own bordered, shaded box, with the
field name, limit, and live character count as a caption ABOVE the box. The
client clicks inside a box, selects all, and pastes into the matching
Psychology Today field, with no risk of grabbing a label or a count.

Delivery: upload the .docx to Google Drive, open it as a Google Doc, share the
link. The boxes survive the conversion as tables.

Usage:
    python3 tools/build_docx.py clients/<slug>-<date>/optimized-profile.md
    python3 tools/build_docx.py <file> -o some/output.docx
    python3 tools/build_docx.py <file> --force   # build even if a block is over limit

The build refuses to run if any block is over its limit or contains an em/en
dash, unless --force is given. Counts in the .docx are recomputed from the copy,
so they are always truthful regardless of what the .md claims.
"""

import argparse
import datetime
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

DASHES = ("—", "–")  # em dash, en dash
LABEL_RE = re.compile(r'"([^"]+)"')
LIMIT_RE = re.compile(r'Limit\s+(\d+)', re.I)
EQ_RE = re.compile(r'^={20,}$')
DASH_RE = re.compile(r'^-{20,}$')

BOX_FILL = "F2F6FB"     # light blue-gray for copy boxes
OVER_COLOR = RGBColor(0xC0, 0x00, 0x00)
MUTED = RGBColor(0x60, 0x60, 0x60)


def parse_doc(text):
    """Walk the deliverable into an ordered list of (kind, data) elements."""
    lines = text.split("\n")
    n, i = len(lines), 0
    out = []

    def blank(s):
        return not s.strip()

    while i < n and blank(lines[i]):
        i += 1
    title = lines[i].strip() if i < n else ""
    i += 1
    subtitle = lines[i].strip() if i < n else ""
    i += 1
    intro = []
    while i < n and not EQ_RE.match(lines[i].strip()) and not DASH_RE.match(lines[i].strip()):
        if lines[i].strip():
            intro.append(lines[i].strip())
        i += 1
    out.append(("title", title))
    out.append(("subtitle", subtitle))
    if intro:
        out.append(("intro", " ".join(intro)))

    while i < n:
        s = lines[i].strip()
        if EQ_RE.match(s):
            i += 1
            banner, closed = [], False
            while i < n:
                if EQ_RE.match(lines[i].strip()):
                    closed = True
                    i += 1
                    break
                if lines[i].strip():
                    banner.append(lines[i].strip())
                i += 1
            joined = " ".join(banner)
            if joined:
                out.append(("section" if closed else "note", joined))
            continue
        if DASH_RE.match(s):
            i += 1
            header = []
            while i < n and not DASH_RE.match(lines[i].strip()):
                header.append(lines[i])
                i += 1
            i += 1  # skip closing divider
            while i < n and blank(lines[i]):
                i += 1
            copy = []
            while i < n and lines[i].strip():
                copy.append(lines[i].strip())
                i += 1
            out.append(("block", build_block(header, " ".join(copy))))
            continue
        if s:
            para = []
            while i < n and lines[i].strip() and not EQ_RE.match(lines[i].strip()) and not DASH_RE.match(lines[i].strip()):
                para.append(lines[i].strip())
                i += 1
            out.append(("note", " ".join(para)))
        else:
            i += 1
    return out


def build_block(header_lines, copy):
    head = "\n".join(header_lines)
    title = header_lines[0].strip() if header_lines else ""
    label_m, limit_m = LABEL_RE.search(head), LIMIT_RE.search(head)
    notes = [ln.strip()[len("Note:"):].strip()
             for ln in header_lines if ln.strip().lower().startswith("note:")]
    return {
        "title": title,
        "label": label_m.group(1) if label_m else "",
        "limit": int(limit_m.group(1)) if limit_m else None,
        "note": " ".join(notes),
        "copy": copy,
        "count": len(copy),
    }


def validate(elements):
    problems = []
    for kind, data in elements:
        if kind != "block":
            continue
        if data["limit"] and data["count"] > data["limit"]:
            problems.append(f'{data["title"]}: over limit ({data["count"]}/{data["limit"]})')
        if any(d in data["copy"] for d in DASHES):
            problems.append(f'{data["title"]}: contains an em/en dash')
    return problems


# --- docx helpers ---------------------------------------------------------

def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def copy_box(doc, copy):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, BOX_FILL)
    para = cell.paragraphs[0]
    run = para.add_run(copy)
    run.font.size = Pt(12)
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    doc.add_paragraph()  # breathing room after the box
    return table


def caption(doc, block):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    field = f'Field: "{block["label"]}"' if block["label"] else "Field"
    over = block["limit"] and block["count"] > block["limit"]
    run = p.add_run(f'{field}   |   {block["count"]} / {block["limit"]} characters')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = OVER_COLOR if over else MUTED
    if block["note"]:
        np = doc.add_paragraph()
        np.paragraph_format.space_after = Pt(2)
        nr = np.add_run(block["note"])
        nr.italic = True
        nr.font.size = Pt(9)
        nr.font.color.rgb = MUTED


def build(md_path, out_path, force=False):
    with open(md_path, encoding="utf-8") as fh:
        text = fh.read()
    elements = parse_doc(text)
    problems = validate(elements)
    if problems and not force:
        print("Refusing to build; fix these (or pass --force):")
        for p in problems:
            print("  - " + p)
        return 1
    if problems:
        print("WARNING building over-limit/dirty copy because --force was given:")
        for p in problems:
            print("  - " + p)

    doc = Document()
    doc.add_paragraph()  # top margin breathing room
    instr = ("How to use this document: each block of copy sits in a shaded box. "
             "Click inside a box, select all of the text in it, and paste it into the "
             "matching Psychology Today field. Copy only what is inside the box; the "
             "labels and character counts above each box are for reference and should "
             "not be pasted.")

    for kind, data in elements:
        if kind == "title":
            h = doc.add_heading(data, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "subtitle":
            p = doc.add_paragraph(data)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.color.rgb = MUTED
        elif kind == "intro":
            ip = doc.add_paragraph()
            ir = ip.add_run(instr)
            ir.italic = True
            ir.font.size = Pt(10)
        elif kind == "section":
            doc.add_heading(data, level=1)
        elif kind == "block":
            doc.add_heading(data["title"], level=2)
            caption(doc, data)
            copy_box(doc, data["copy"])
        elif kind == "note":
            np = doc.add_paragraph()
            nr = np.add_run(data)
            nr.font.size = Pt(9)
            nr.font.color.rgb = MUTED

    doc.save(out_path)
    n_blocks = sum(1 for k, _ in elements if k == "block")
    print(f"Built {out_path} ({n_blocks} copy blocks).")
    print("Next: upload to Google Drive, open as a Google Doc, share the link.")
    return 0


def default_output(md_path):
    """Deliverable lands in the client folder, namespaced with the client slug
    and a datetimestamp so revision runs never overwrite an earlier delivery."""
    folder = os.path.dirname(os.path.abspath(md_path))
    base = os.path.basename(folder)
    # client folders are <slug>-<YYYY-MM-DD>; strip the trailing date to get the slug
    m = re.match(r"^(.*)-\d{4}-\d{2}-\d{2}$", base)
    slug = m.group(1) if m else base
    stamp = datetime.datetime.now().strftime("%Y-%m-%d-%H%M%S")
    return os.path.join(folder, f"{slug}-optimized-profile-{stamp}.docx")


def main():
    ap = argparse.ArgumentParser(description="Build a client-friendly .docx from optimized-profile.md.")
    ap.add_argument("file", help="path to optimized-profile.md")
    ap.add_argument("-o", "--out", help="output .docx path (default: client folder, namespaced with slug + datetimestamp)")
    ap.add_argument("--force", action="store_true", help="build even if a block is over limit or has a dash")
    args = ap.parse_args()
    if not os.path.isfile(args.file):
        print(f"File not found: {args.file}")
        sys.exit(2)
    out = args.out or default_output(args.file)
    sys.exit(build(args.file, out, force=args.force))


if __name__ == "__main__":
    main()
